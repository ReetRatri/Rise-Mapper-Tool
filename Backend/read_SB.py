# file  -- exportContent.py --
import time
import docx
#import io
import os
import json
import re
import server

#import UI_layout
#import textract
# import PySimpleGUI as sg
import os, shutil

from pathlib import Path

import traceback

global downloads_path
global sbPageCount
global page_data

global course_page_no
global specific_text
global cell_contents
cell_contents=[]

sbPageCount = []
page_data=[]
downloads_path = str(Path.home() / "Downloads")
downloads_path = downloads_path.replace('\\','/')
downloads_path = downloads_path+"/Comparison_Report/"


os.makedirs(downloads_path, exist_ok=True)

#global slideCounttotal
global SB_mydoc
# global docx_file_path


def find_table_with_ost(docx_file):
    # Load the DOCX file
    doc = docx.Document(docx_file)


    # Iterate through each table in the document
    for table_index, paragraph in enumerate(doc.paragraphs):

    # for table_index, table in enumerate(doc):

        # Iterate through each cell in the table
        # for row in table.rows:
        #     for cell in row.cells:
        cell=paragraph.text
        # print('its working')
                # Check if "OST" is in the cell text
        if "OST:" in cell:
            return table_index  # Return the index of the table

    # If "OST" is not found in any table, return None
    return None


def deleteFiles():

    folder = downloads_path
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            UI_layout.sg.Popup('Failed to delete %s. Reason: %s' % (file_path, e))


# Function to find the table index for a given Slide Title




SB_mydoc = docx.Document()



global sbdata_text
sbdata_text = []
def find_table_with_specific_text(docx_file, specific_text,parentDataText):
    # Open the DOCX file
    #doc = Document(docx_file)
    doc = docx.Document(docx_file)
    sbdata_text = ""
    sbdata_text2 = ""
    for paragraph in doc.paragraphs:
    # Use regex to remove trailing tab followed by a digit
        # cleaned_text = re.sub(r'\t\d$', '', paragraph.text.strip())
        # sbdata_text += cleaned_text + "\n"
        # Strip leading and trailing whitespace, including newlines
        cleaned_text = paragraph.text.strip()
        # Use regex to remove trailing tab followed by a digit
        cleaned_text = re.sub(r'\t\d$', '', cleaned_text)
        # Append the cleaned text with a newline
        sbdata_text += cleaned_text + "\n"

       
    print('---------------------------------------------------------------------------------------------')
    print('---------------------------------------------------------------------------------------------')
    print('---------------------------------------------------------------------------------------------')
    print(parentDataText)
    # print(sbdata_text)
    pattern = r"/\*+Page+ - start \*+\s*(.*?)\s*/\*+Page+ - end \*+"

    # page_data = []
    page_start_split = re.split(r'/\*+Page\s*-\s*start\s*\*+/', sbdata_text)
    for content in page_start_split[1:]:
        # content_between = content.split("/********Page - end ********/")[0].strip()
        page_end_delimiter = "/\*+Page\s*-\s*end\s*\*+/"
        content_between = re.split(page_end_delimiter, content)[0].strip()
        # print(content_between)
        page_data.append(content_between)

    # page_start_split2 = re.split(r'/\*+Page\s*-\s*start\s*\*+/', sbdata_text2)
    # for content in page_start_split2[1:]:
    #     page_end_delimiter2 = "/\*+Page\s*-\s*end\s*\*+/"
    #     content_between2 = re.split(page_end_delimiter2, content)[0].strip()
    #     for i, item in enumerate(page_data):
    #         if specific_text in page_data[i] and  specific_text in content_between2:
    #             page_data[i] + '/n' + '/n' + content_between2
    #         print(page_data[i])




    # if specific_text in page_data:

        # for i, item in enumerate(page_data):


    # Find all matches of the pattern in the document
    # matches = re.findall(pattern, sbdata_text, re.DOTALL)
    # global page_data
    # page_data = []
    # Extracted page data will be stored in a list

    # if specific_text =="Course overview" or specific_text=="Quiz" :
    #     pass
    # else:
    #     course_page_no=int(specific_text.split(' ')[1])


    # Iterate through the matches and collect page data
    # for match in matches:
    #     page_data.append(match.strip())

    if specific_text!="Course overview" and specific_text!="Quiz" and specific_text!="Final assessment":

        for i, item in enumerate(page_data):
            if specific_text in item:
                cell_contents=page_data[i]
                break

    elif specific_text=="Quiz" or specific_text=="Final assessment" :
            print('true')
            for i, item in enumerate(page_data):
                if specific_text in item:
                    if parentDataText in item:
                        cell_contents=page_data[i]
                        break
                else:
                    print('something missing  in sb')



    else:
        # a= page_data[0]
        # pattern2=re.compile(r"/\*+Page\d+ - end\*+")
        # parts = re.split(pattern2, a)

        if(specific_text in  page_data[0]):
            cell_contents=page_data[0]
        else:
            cell_contents="Please check there  is  some issue "


# reet

    return cell_contents,page_data  # Return None if the table is not found






# Function to find the content from each cell of a specific table
def find_content_in_table(page_data):
    # Open the DOCX file

    # doc = docx.Document(docx_file)

    # Initialize a list to store the content from each cell
    cell_contents = []




    # if 'Page Type' in cell_text:
    #     pageType = cell_text.split('Page Type:')[1].strip()
    #     if pageType == 'Video':
    #         #f.write("this is video screen")
    #         cell_contents.append('this is video page')
            # Check if the cell does not contain "OST:"
# reet
    # if(specific_text in  page_data[int(specific_text.split(' ')[1])]):

    #     cell_contents.append((page_data[int(specific_text.split(' ')[1])]))

    # if('Course overview' in page_data[0]):
    #     cell_contents.append((page_data[0]))

    if specific_text!="Course overview" and specific_text!="Quiz" and specific_text!="Final assessment":
        course_page_no=int(specific_text.split(' ')[1])
        # if course_page_no ==1:
        #     a= page_data[0]
        #     pattern2=re.compile(r"/\*+Page\d+ - end\*+")
        #     parts = re.split(pattern2, a)
        #     pattern3 =re.compile(r"/\*+Page\d+ - start \*+")
        #     part=re.sub(pattern3, '', parts[1])
        #     cell_contents.append(part)



        # else:
        for i, item in enumerate(page_data):
            if specific_text in item:
                    # cell_contents=page_data[i]
                cell_contents.append(page_data[i])
                break
            # cell_contents.append(page_data[int(course_page_no)-1])
    elif specific_text=="Quiz" or specific_text=="Final assessment":
            print('true')
            for i, item in enumerate(page_data):
                if specific_text in item:
                    parentDataText=''
                    if parentDataText in item:
                        cell_contents.append(page_data[i])
                        break
                else:
                    print('something missing  in sb')
    else:
        # a= page_data[0]
        # pattern2=re.compile(r"/\*+Page\d+ - end\*+")
        # parts = re.split(pattern2, a)

        if(specific_text in  page_data[0]):
            cell_contents.append(page_data[0])
        else:
            a="Please check there  is  some issue "
            cell_contents.append(a)







# reet

    # if 'OST' in cell_text:
    #     if 'ON-SCREEN TEXT (OST):' in cell_text:
    #         content = cell_text.split('OST:')[1].strip()
    #         if content not in cell_contents:
    #             cell_contents.append(content)
    #         #f.write((cel.text.split('ON-SCREEN TEXT (OST):')[1]).strip() + '\n')
    #         cell_text = ''
    #     elif 'OST:' in cell_text:
    #         content = cell_text.split('OST:')[1].strip()
    #         if content not in cell_contents:
    #             cell_contents.append(content)
    #         #f.write((cel.text.split('OST:')[1]).strip() + '\n')
    #         cell_text = ''


    return cell_contents

# Function to write content to a text file
def write_content_to_txt(cell_contents, txt_file,parentDataText):
    # if cell_contents ==text = re.sub(r'(/\n{3,}/)+', '\n', text)
    # pattern= re.compile(r'(/\n{3,}/)+')
    # cell_contents = re.sub(r'(/\n{3,}/)+', '\n', cell_contents)
    print(type(cell_contents))
    print(parentDataText)


    with open(txt_file, "w", encoding="utf-8") as file:
        file.write('Storyboard Text'+'\n\n')
        for content in cell_contents:
            content= re.sub(r'(/\n{3,}/)+', '\n', content)
            file.write(content + "\n")

def readSB(specific_text, SB_path,parentDataText):

    # Specify the DOCX file path and specific text to check for
    docx_file_path = SB_path
    specific_text_to_check = specific_text
    cell_contents=[]
    # sg.Popup(slideTitle)
    # Call the function to find the table with the specific text
    print(parentDataText)
    table_index = find_table_with_specific_text(docx_file_path, specific_text_to_check,parentDataText)

    # Check if a table with the specific text was found
    # if page_data is not []:
    if page_data is not []:
        # print(f"Table containing '{specific_text_to_check}' found at index {table_index}")
        file_path = os.path.join(server.app.config['UPLOAD_FOLDER'])
        output_txt_file = file_path+'/Storyboard_'+specific_text_to_check+'.txt'
        print('-----'+ specific_text_to_check)

        # Call the function to find content from each cell of the specified table
        # cell_contents = find_content_in_table(page_data)
        # if specific_text !='Course overview':

        #     if(specific_text in  page_data[int(specific_text.split(' ')[1])-1]):
        #         cell_contents=[(page_data[int(specific_text.split(' ')[1])-1])]
        # else:
        #     if('Course overview' in page_data[0]):
        #         cell_contents.append((page_data[0]))



        # reet


        if specific_text!="Course overview" and specific_text!="Quiz" and specific_text!="Final assessment":
            course_page_no=int(specific_text_to_check.split(' ')[1])

            # else:
            for i, item in enumerate(page_data):
                if specific_text in item:
                    # cell_contents=page_data[i]
                    cell_contents.append(page_data[i])
                    break

                # cell_contents.append(page_data[int(course_page_no)-1])
        elif specific_text=="Quiz" or specific_text=="Final assessment":
            print('true')
            for i, item in enumerate(page_data):
                if specific_text in item:
                    if parentDataText in item:
                        cell_contents.append(page_data[i])
                        break
        else:
            # a= page_data[0]
            # pattern2=re.compile(r"/\*+Page\d+ - end\*+")
            # parts = re.split(pattern2, a)

            if(specific_text in  page_data[0]):
                cell_contents.append(page_data[0])
            else:
                a="Please check there  is  some issue "
                cell_contents.append(a)

        # reet

        # Write the content to the specified text file
        # cell_contents=page_data[(int(specific_text.split(' ')[1]))-1]
        # cell_contents.append(page_data[int(specific_text.split(' ')[1])])
        write_content_to_txt(cell_contents, output_txt_file,parentDataText)
        output_txt_file_path = os.path.abspath(output_txt_file)
        return output_txt_file_path
    else:
        print(f"No table containing '{specific_text_to_check}' found in the document")
        # return
    print('+++++++++'+course_page_no)

    # table_index_with_ost = find_table_with_ost(SB_path)

    try:
        # doc = docx.Document(SB_path)

        # tablelist = []

        SB_mydoc.add_paragraph('Storyboard Text'+'\n\n')
        file_path = os.path.join(server.app.config['UPLOAD_FOLDER'])
        with open(file_path+'/test.txt', 'w') as test:
            test.write(SB_path)


            with open(file_path+'/Storyboard_'+course_page_no+'.txt', 'w') as f:
                f.write('Storyboard Text'+'\n\n')


                table_index = None







                f.write(cell_contents)
                # cel.text = ''



        SB_mydoc.save(downloads_path+"SB_Content.docx")
    except Exception as e:
        print(f'issue in lesscss{e}')
