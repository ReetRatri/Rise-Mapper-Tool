#file  -- exportContent.py --
import time
from xmlrpc.client import boolean
import docx
#import io
import os
import json
import difference


# import re


import pages

from docx.shared import Inches

import read_SB

global routPath

global storyMyDoc

storyMyDoc = docx.Document()



def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def removeBlankRows(doc):
    for table in doc.tables:
        for row in table.rows:
            for j, cel in enumerate(row.cells):
                if cel.text == '':
                    remove(cel)
    return doc

def remove(self):
    self._parent._tbl.remove(self._tr)

def createTable(item, mydoc,val,lable):

    table2 = mydoc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    row = table2.add_row().cells
    row[0].text = lable
    row[1].text = val
    #make_rows_bold(table1.rows[0])
    #make_rows_bold(table2.rows[0])
    #print(os.path.join(root_path, item))
    #time.sleep(0.05)
    #UI_layout.progress_bar.UpdateBar(counterr_File2, UI_layout.increase_File_Count2) #pylint:disable=[progress_bar.UpdateBar(j + 1, increase_File_Count2)]

# Create your dictionary class
class my_dictionary(dict):

  # __init__ function
  def __init__(self):
    self = dict()

  # Function to add key:value
  def add(self, key, value):
    self[key] = value
# Main Function
#dict_obj = my_dictionary()
global dataSource

dataSource = my_dictionary()
global masterData
masterData = []
def exportC(xmlFiles):
    #print(UI_layout.increase_File_Count2)



    try:
        routPath = read_SB.downloads_path

        os.makedirs(routPath, exist_ok=True)
        # setPages(xmlFiles)
        counterr_File2 = 0
        # mydoc = docx.Document()
        # #print(xmlFiles)
        # document_Texts = []
        res = []
        global val_4_test
        counterdata = 0
        for root_path, directories, files in xmlFiles:
            for idx, val in enumerate(pages.pages):
                # print('root_path,val')
                str1 = ''
                slideData = []
                val_4_test=val
                path = root_path +'/'+ val
                with open(path, 'r', encoding='utf-8', errors='ignore') as file:
                    for line in file:
                        js_data = str1.join(line)
                        js_data1 = js_data.replace("window.globalProvideData('slide', '",'')
                        js_data1 = js_data1.replace("\ufeff",'')
                        js_data1 = js_data1[:-3]
                        js_data1 = js_data1.replace('\\\\"','\\"')
                        js_data1 = js_data1.replace("\\\'","\'")
                        js_data1 = js_data1.split('"text":"')
                        docName = 'slide_'+str(idx+1)
                        dataSource[docName]= []
                        for c1, a_ in enumerate(js_data1):
                            if c1 != 0:
                                if '",\"style\":' in a_:
                                    text = a_.split('",\"style\":')[0]
                                    dataSource[docName].append(text.replace("\\n", "").replace("%([^%]+)%","").replace("%_player.page%","").replace("%_playerVars.menuSlideTitle%","").replace('"', "").replace('\\r', "\n").replace("\\","").strip())
                                elif '","style":' in a_:
                                    text = a_.split('","style":')[0]
                                    dataSource[docName].append(text.replace("\\n", "").replace("%([^%]+)%","").replace("%_player.page%","").replace("%_playerVars.menuSlideTitle%","").replace('"', "").replace('\\r', "\n").replace("\\","").strip())

            shortData(dataSource)

           # mydoc.save("text.docx")
    except:
        pass
        #UI_layout.sg.Print(f'An error happened.  Here is the info:')
        #UI_layout.sg.popup_ok('Please validate the Image -- ',os.path.join(root_path, item),text_color = 'Red',title = 'Error',background_color ='#dadce5')
        # print(f'issue in lesscss')
        #print('')
        # UI_layout.sg.Popup('Please validate.')

def shortData(data):
    #valueList = list(data.values())
    keyList = list(data.keys())
    for k, v in enumerate(keyList):
        #print(list(data[v].values()))
        #valKeyList = list(data[v].keys())
        for k_, v_ in enumerate(data[v]):
            # print(data[v][v_])
            #if data[v] != '':
            checkAvailability(data[v], v, v_)
        #
        # dataKey = '***objects***_'+str(k+1)
        # data[v][dataKey]
    for a in dataSource:
        dataSource[a] = [i for i in dataSource[a] if i]
        dataSource[a] = [i for i in [*set(dataSource[a])] if i]


    prep_Doc_(dataSource)
# reet
# def checkAvailability(val, cuurentKey,currentValue):
#     # for ele in dataSource:
#     #     if (currentValue=='	%_playerVars.menuSlideTitle%'):

#     # return
#     dataCount = {}
#     global extrawords
#     for ele in dataSource:

#         if cuurentKey != ele:
#             for k , v in enumerate(dataSource[ele]):
#                 if v == currentValue:
#                     pass
#                     for values in dataSource.values():
#                         for value in values:
#                             if value in dataCount:
#                                 dataCount[value] += 1
#                             else:
#                                 dataCount[value] = 1
#                     # print(dataCount)
#                     extrawords= {key: value for key, value in dataCount.items() if value >=(len(ele))}
#                     extrawords=set(extrawords.keys())
#                     extrawords.discard('')


                    # if dataSource[ele][k] not in masterData and dataSource[ele][k] != '':
                    #     masterData.append( dataSource[ele][k])
                    # dataSource[ele][k] =''

def checkAvailability(val, cuurentKey,currentValue):
    # for ele in dataSource:
    #     if (currentValue=='	%_playerVars.menuSlideTitle%'):

    # return

    for ele in dataSource:

        if cuurentKey != ele:
            for k , v in enumerate(dataSource[ele]):
                if v == currentValue:
                    # if functools.reduce(lambda i, j : i and j, map(lambda m, k: m == k, val, dataSource[ele][v]), True) :
                    #     sParentKey_ = ele
                    #     sobjKey_ = v
                        #if dataSource[ele][v] != '':
                        #if sParentKey_ != parentKey:
                        # if sParentKey_ != parentKey and cuurentKey != sobjKey_:
                            # print('identical')
                    if dataSource[ele][k] not in masterData and dataSource[ele][k] != '':
                        masterData.append( dataSource[ele][k])
                    # dataSource[ele][k] =''
                    # elif sorted(val) == sorted(dataSource[ele][v]):
                    #     sParentKey_ = ele
                    #     sobjKey_ = v
                    #     #if dataSource[ele][v] != '':
                    #     # if sParentKey_ != parentKey and cuurentKey != sobjKey_:
                    #     #if sParentKey_ != parentKey:
                    #         # print('identical')
                    #     dataSource[sParentKey_][sobjKey_] =''

def checkMasterData(v):
    for i , el in enumerate(v):
        key = 'slide_'+ str(i+1)
        for ind, val in enumerate(v[key]):
            if el == val:
                v[key][ind] = ''

def prep_Doc_(v):
    checkMasterData(v)
    encountered_values = set()
    masterData.append("")
    paragraph = storyMyDoc.add_paragraph('Storyline Content',style = 'Heading 1').add_run()
    for ele in v:
        paragraph = storyMyDoc.add_paragraph(ele,style = 'Heading 1').add_run()
        with open(read_SB.downloads_path+'Storyline_'+ele+'.txt', 'x') as f:
            f.write('Storyline Content'+'\n\n')
            f.write('\n\n'+ ele +'\n\n')
            if diff.video_check==True:
                f.write('\n\n'+ "   VIDEO FILE" +'\n\n')


            for k_ , v_ in enumerate(v[ele]):
                if v_ not in masterData and v_ not in encountered_values:
                    # if v_ != '' and v_ not in masterData:
                    encountered_values.add(v_)

                    paragraph = storyMyDoc.add_paragraph().add_run(v_)
                    f.write(v_+'\n')
    storyMyDoc.save(read_SB.downloads_path+"Storyline_text.docx")

